"""
Read all files from xuejie_files, extract text, structure into Q&A pairs,
and write to database.db with auto-deduplication.
"""
import os, re, sqlite3, sys
from io import StringIO

# ── File paths ──
DATA_DIR = os.path.join(os.path.dirname(__file__), '..', 'xuejie_files', 'xuejie_files')
DB_PATH = os.path.join(os.path.dirname(__file__), 'database.db')

# ── Existing categories from app.py ──
CATEGORIES = ["图书馆", "食堂", "宿舍", "校园服务", "选课", "考试", "毕业要求", "奖学金", "德育积分", "综合素质实践", "社团", "其他"]

# ── PDF extraction ──
def extract_pdf(path):
    from pdfminer.high_level import extract_text
    return extract_text(path)

# ── DOCX extraction ──
def extract_docx(path):
    from docx import Document
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs]
    # Also extract tables
    rows = []
    for table in doc.tables:
        for row in table.rows:
            rows.append(' | '.join(cell.text.strip() for cell in row.cells))
    text = '\n'.join(paras)
    if rows:
        text += '\n\n' + '\n'.join(rows)
    return text

# ── XLSX extraction ──
def extract_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    lines = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        lines.append(f'=== {sheet} ===')
        for row in ws.iter_rows(values_only=True):
            vals = [str(v).strip() if v else '' for v in row]
            if any(v for v in vals):
                lines.append(' | '.join(vals))
    return '\n'.join(lines)

# ── Extract text by file type ──
def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext == '.pdf':
            return extract_pdf(filepath)
        elif ext in ('.docx', '.doc'):
            return extract_docx(filepath)
        elif ext in ('.xlsx', '.xls'):
            return extract_xlsx(filepath)
        elif ext == '.txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return ''
    except Exception as e:
        return f'[ERROR extracting {filepath}: {e}]'

# ── Categorize file ──
def categorize_file(filename):
    """Map filename to a category based on its content."""
    name = filename.lower()
    if '图书馆' in name or '入馆' in name:
        return '图书馆'
    if '食堂' in name:
        return '食堂'
    if '宿舍' in name:
        return '宿舍'
    if '选课' in name:
        return '选课'
    if '考试' in name:
        return '考试'
    if '毕业' in name:
        return '毕业要求'
    if '奖学金' in name or '奖助' in name:
        return '奖学金'
    if '德育' in name:
        return '德育积分'
    if '综合素质' in name or '实践' in name:
        return '综合素质实践'
    if '社团' in name:
        return '社团'
    if '校园' in name or '服务' in name or '校园卡' in name or '保险' in name or '困难认定' in name or '思政' in name or '学平险' in name:
        return '校园服务'
    if '竞赛' in name or '学科竞赛' in name:
        return '综合素质实践'
    if '岗位' in name or '固定岗' in name:
        return '综合素质实践'
    if '学生手册' in name:
        return None  # multi-category, will be processed separately
    return '其他'

# ── Parse structured Q&A from text ──
def parse_qa_from_text(text, category, filename):
    """
    Given extracted text and category, produce list of (question, answer) tuples.
    Uses heuristics to split text into meaningful Q&A chunks.
    """
    pairs = []
    lines = text.split('\n')

    # Remove empty lines and strip
    lines = [l.strip() for l in lines if l.strip()]

    if not lines:
        return pairs

    # For DOCX/XLSX with table-like content, each "line" might be a Q&A
    # Heuristic: lines with | separators might be table rows

    # Try to extract based on content structure
    # Strategy: look for section headers and pair them with content

    # First, try table-based extraction (lines with | or \t)
    table_rows = [l for l in lines if '|' in l or '\t' in l]
    if len(table_rows) > 5:
        for row in table_rows:
            cols = re.split(r'\s*\|\s*|\t+', row)
            cols = [c.strip() for c in cols if c.strip()]
            if len(cols) >= 2:
                # First col = question, rest = answer
                q, a = cols[0], ' '.join(cols[1:])
                if len(q) > 3 and len(a) > 3:
                    pairs.append((q, a))

    # If table extraction didn't yield much, try topic-based extraction
    if len(pairs) < 3:
        pairs = []
        current_q = None
        current_a_lines = []

        for line in lines:
            # Detect potential question/header lines
            # Ends with ? or ？, or starts with number+。
            # or is a section header in 【】or （）or 【】
            is_header = False
            if re.match(r'^[\d一二三四五六七八九十]+[\.、．\)）]', line):
                is_header = True
            if '?' in line or '？' in line:
                is_header = True
            if re.match(r'^[\[【（(][^】）)\]]+[\]】）)]', line):
                is_header = True
            if line.endswith('：') or line.endswith(':') or line.endswith('如下') or line.endswith('办法') or line.endswith('规定') or line.endswith('细则') or line.endswith('说明') or line.endswith('介绍'):
                is_header = True

            # Filter out very long lines as content, not headers
            if is_header and len(line) < 100:
                if current_q and current_a_lines:
                    a_text = ''.join(current_a_lines).strip()
                    if len(a_text) > 5:
                        pairs.append((current_q, a_text))
                current_q = line
                current_a_lines = []
            elif current_q:
                current_a_lines.append(line)

        if current_q and current_a_lines:
            a_text = ''.join(current_a_lines).strip()
            if len(a_text) > 5:
                pairs.append((current_q, a_text))

    # If still no pairs, just make the whole text an answer
    if not pairs and len(''.join(lines)) > 20:
        # Create one big entry
        title = os.path.splitext(filename)[0]
        pairs.append((f'{title} 相关内容', '\n'.join(lines[:50])))

    return pairs

# ── Smart extraction for student handbook ──
def parse_student_handbook(text):
    """Parse the student handbook into multiple categories."""
    pairs_by_cat = {}
    current_cat = None
    current_q = None
    current_a = []

    # Category keywords to detect sections
    cat_keywords = {
        '选课': ['选课', '课程', '学分'],
        '考试': ['考试', '考核', '成绩', '考场'],
        '毕业要求': ['毕业', '学位', '结业', '肄业'],
        '奖学金': ['奖学金', '奖励', '表彰'],
        '德育积分': ['德育', '处分', '违纪', '奖励'],
        '宿舍': ['宿舍', '公寓', '住宿'],
        '图书馆': ['图书馆', '借阅'],
        '食堂': ['食堂', '餐饮', '就餐'],
        '校园服务': ['校园卡', '网络', '服务', '医疗', '资助', '贷款'],
        '综合素质实践': ['实践', '创新', '竞赛', '科研'],
        '社团': ['社团', '组织', '活动'],
    }

    lines = [l.strip() for l in text.split('\n') if l.strip()]

    for line in lines:
        # Detect chapter/section
        for cat, keywords in cat_keywords.items():
            if any(kw in line for kw in keywords):
                current_cat = cat
                break

        # Detect questions within sections
        if '?' in line or '？' in line or re.match(r'^[\d]+[\.、．]', line):
            if current_q and current_a:
                a_text = ' '.join(current_a).strip()
                if len(a_text) > 5:
                    cat = current_cat or '其他'
                    if cat not in pairs_by_cat:
                        pairs_by_cat[cat] = []
                    pairs_by_cat[cat].append((current_q, a_text))
            current_q = line
            current_a = []
        elif current_q:
            current_a.append(line)

    # Last Q
    if current_q and current_a:
        a_text = ' '.join(current_a).strip()
        if len(a_text) > 5:
            cat = current_cat or '其他'
            if cat not in pairs_by_cat:
                pairs_by_cat[cat] = []
            pairs_by_cat[cat].append((current_q, a_text))

    return pairs_by_cat

# ── Write to database ──
def write_to_db(all_pairs):
    """Write all Q&A pairs to database, deduplicating by question text."""
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    total_new = 0
    total_dup = 0

    for cat, pairs in all_pairs.items():
        if cat not in CATEGORIES:
            cat = '其他'

        # Ensure table exists
        cur.execute(f'CREATE TABLE IF NOT EXISTS "{cat}" (id INTEGER PRIMARY KEY AUTOINCREMENT, question TEXT, answer TEXT)')

        # Get existing questions
        existing = set()
        try:
            rows = cur.execute(f'SELECT question FROM "{cat}"').fetchall()
            existing = set(r[0] for r in rows)
        except:
            pass

        for q, a in pairs:
            q = q.strip()
            a = a.strip()
            if not q or not a:
                continue
            if len(q) < 4 or len(a) < 4:
                continue  # too short, skip

            # Dedup by question text
            if q in existing:
                total_dup += 1
                continue

            cur.execute(f'INSERT INTO "{cat}" (question, answer) VALUES (?, ?)', (q, a))
            existing.add(q)
            total_new += 1

        conn.commit()

    conn.close()
    return total_new, total_dup

# ══════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════
def main():
    all_pairs = {}  # category -> [(q, a), ...]

    files = sorted(os.listdir(DATA_DIR))
    print(f"Found {len(files)} files")

    for fname in files:
        fpath = os.path.join(DATA_DIR, fname)
        if not os.path.isfile(fpath):
            continue

        print(f"\n{'='*60}")
        print(f"Processing: {fname}")
        print(f"{'='*60}")

        text = extract_text(fpath)
        text_len = len(text)
        print(f"  Extracted {text_len} chars")

        if text_len < 10:
            print(f"  SKIP: too short")
            continue

        category = categorize_file(fname)

        if fname == '2023版学生手册（本科生）.pdf':
            # Student handbook: multi-category
            handbook_pairs = parse_student_handbook(text)
            for cat, pairs in handbook_pairs.items():
                if cat not in all_pairs:
                    all_pairs[cat] = []
                all_pairs[cat].extend(pairs)
                print(f"  Handbook -> {cat}: {len(pairs)} Q&A")
        elif category:
            # Single category file
            pairs = parse_qa_from_text(text, category, fname)
            if pairs:
                if category not in all_pairs:
                    all_pairs[category] = []
                all_pairs[category].extend(pairs)
                print(f"  → {category}: {len(pairs)} Q&A extracted")
            else:
                print(f"  WARNING: no Q&A extracted for {category}")
                # Fallback: save whole text as one entry
                title = os.path.splitext(fname)[0]
                if category not in all_pairs:
                    all_pairs[category] = []
                all_pairs[category].append((f'{title} 相关内容', text[:500]))
                print(f"  → {category}: 1 fallback entry")
        else:
            print(f"  SKIP: could not determine category")

    # Summary
    print(f"\n{'='*60}")
    print("EXTRACTION SUMMARY")
    print(f"{'='*60}")
    for cat in CATEGORIES:
        if cat in all_pairs:
            print(f"  {cat}: {len(all_pairs[cat])} pairs")

    total = sum(len(v) for v in all_pairs.values())
    print(f"\nTotal Q&A pairs to insert: {total}")

    # Write to DB
    if total > 0:
        new, dup = write_to_db(all_pairs)
        print(f"\nDatabase write: {new} new, {dup} duplicates skipped")
        print(f"Done! Check your data in database.db")
    else:
        print("\nNo data to write.")

if __name__ == '__main__':
    main()
