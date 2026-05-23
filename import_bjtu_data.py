"""
Import BJTU training plan PDFs from Desktop/bjtu into database.db
"""
import os, re, sqlite3
from pdfminer.high_level import extract_text
from docx import Document

DATA_DIR = os.path.expanduser('~/Desktop/bjtu')
DB_PATH = os.path.join(os.path.dirname(__file__), 'database.db')

CATEGORIES = ["图书馆", "食堂", "宿舍", "校园服务", "选课", "考试", "毕业要求", "奖学金", "德育积分", "综合素质实践", "社团", "其他", "专业培养方案"]

# ── Extract text ──
def extract_pdf(path):
    return extract_text(path)

def extract_docx(path):
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs]
    rows = []
    for table in doc.tables:
        for row in table.rows:
            rows.append(' | '.join(cell.text.strip() for cell in row.cells))
    text = '\n'.join(paras)
    if rows:
        text += '\n\n' + '\n'.join(rows)
    return text

def extract_text_file(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def extract_text(fp):
    ext = os.path.splitext(fp)[1].lower()
    if ext == '.pdf': return extract_pdf(fp)
    elif ext == '.docx': return extract_docx(fp)
    elif ext == '.txt': return extract_text_file(fp)
    return ''

# ── Parse text into Q&A pairs ──
def extract_school_intro(text, filename):
    """Extract school introduction info."""
    pairs = []
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return pairs

    # Try to find school name from first few lines
    school_name = os.path.splitext(filename)[0]
    title_text = ' '.join(lines[:5])

    # School introduction - look for 介绍/概况/目录
    intro_lines = []
    in_intro = False
    for line in lines:
        if '介绍' in line or '概况' in line:
            in_intro = True
        if '目录' in line or '培养计划' in line or '课程' in line:
            in_intro = False
        if in_intro:
            intro_lines.append(line)

    if intro_lines:
        intro = ''.join(intro_lines[:10])
        if len(intro) > 20:
            pairs.append((f'{school_name}的情况如何？', intro[:500]))

    # Try to extract Q&A from content
    current_q = None
    current_a = []

    for line in lines:
        is_q = False
        if re.match(r'^[\d一二三四五六七八九十]+[\.、．\)）]', line):
            is_q = True
        if '?' in line or '？' in line:
            is_q = True
        if line.endswith('：') or line.endswith(':'):
            is_q = True

        if is_q and len(line) < 100:
            if current_q and current_a:
                a = ''.join(current_a).strip()
                if len(a) > 10:
                    pairs.append((current_q, a))
            current_q = line
            current_a = []
        elif current_q:
            current_a.append(line)

    if current_q and current_a:
        a = ''.join(current_a).strip()
        if len(a) > 10:
            pairs.append((current_q, a))

    if not pairs:
        pairs.append((f'{school_name} 相关信息', ''.join(lines[:30])[:500]))

    return pairs

# ── Categorize by filename ──
def categorize(filename):
    name = filename.lower()
    if '交通运输' in name: return '其他'
    if '经管' in name or '经济' in name: return '其他'
    if '计算机' in name: return '其他'
    if '自动化' in name: return '其他'
    if '电子' in name or '通信' in name or '信息' in name: return '其他'
    if '培养方案' in name: return '其他'
    if '双学' in name: return '其他'
    return '其他'

# ── Write to DB ──
def write_to_db(all_pairs):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    cat_name = '专业培养方案'
    cur.execute(f'''CREATE TABLE IF NOT EXISTS "{cat_name}" (
        id INTEGER PRIMARY KEY AUTOINCREMENT, question TEXT, answer TEXT
    )''')

    existing = set()
    try:
        rows = cur.execute(f'SELECT question FROM "{cat_name}"').fetchall()
        existing = set(r[0] for r in rows)
    except: pass

    new = 0
    dup = 0
    for cat, pairs in all_pairs.items():
        for q, a in pairs:
            q = q.strip()[:200]
            a = a.strip()[:2000]
            if not q or not a or len(q) < 4 or len(a) < 4:
                continue
            if q in existing:
                dup += 1
                continue
            cur.execute(f'INSERT INTO "{cat_name}" (question, answer) VALUES (?, ?)', (q, a))
            existing.add(q)
            new += 1
    conn.commit()
    conn.close()
    return new, dup

# ── MAIN ──
def main():
    all_pairs = {}
    files = sorted(os.listdir(DATA_DIR))
    print(f"Found {len(files)} files in {DATA_DIR}")

    for fname in files:
        fpath = os.path.join(DATA_DIR, fname)
        if not os.path.isfile(fpath):
            continue

        print(f"\nProcessing: {fname}")

        text = extract_text(fpath)
        if len(text) < 20:
            print(f"  SKIP: too short ({len(text)} chars)")

        pairs = extract_school_intro(text, fname)
        cat = categorize(fname)
        if cat not in all_pairs:
            all_pairs[cat] = []
        all_pairs[cat].extend(pairs)
        print(f"  → {len(pairs)} Q&A pairs extracted ({len(text)} chars)")

    total = sum(len(v) for v in all_pairs.values())
    print(f"\nTotal Q&A: {total}")
    if total > 0:
        new, dup = write_to_db(all_pairs)
        print(f"DB: {new} new, {dup} duplicates skipped")
    else:
        print("No data to write.")

if __name__ == '__main__':
    main()
